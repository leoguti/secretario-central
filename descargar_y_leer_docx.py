#!/usr/bin/env python3
"""
Script para descargar y leer archivo Word desde Shared Drive
"""

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import json
import io

DOC_ID = '1z0HYH29b_JHrEPIMQuZD73Kt5cO7n4WX'
TOKEN_PATH = 'token_trufi.json'
OUTPUT_FILE = '/tmp/anexo_tecnico_trujillo.docx'

def get_drive_service():
    with open(TOKEN_PATH, 'r') as f:
        token_data = json.load(f)

    creds = Credentials(
        token=token_data['token'],
        refresh_token=token_data['refresh_token'],
        token_uri=token_data['token_uri'],
        client_id=token_data['client_id'],
        client_secret=token_data['client_secret'],
        scopes=token_data['scopes']
    )

    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        token_data['token'] = creds.token
        with open(TOKEN_PATH, 'w') as f:
            json.dump(token_data, f, indent=2)

    return build('drive', 'v3', credentials=creds)

def main():
    service = get_drive_service()

    print("Descargando archivo Word...")

    # Descargar archivo
    request = service.files().get_media(fileId=DOC_ID)
    file = io.BytesIO()
    downloader = MediaIoBaseDownload(file, request)

    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print(f"Descarga: {int(status.progress() * 100)}%")

    # Guardar archivo
    with open(OUTPUT_FILE, 'wb') as f:
        f.write(file.getvalue())

    print(f"✅ Archivo guardado en: {OUTPUT_FILE}")

    # Leer con python-docx
    try:
        from docx import Document

        doc = Document(OUTPUT_FILE)

        print("\n" + "=" * 60)
        print("CONTENIDO DEL DOCUMENTO")
        print("=" * 60 + "\n")

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detectar títulos por estilo
                if para.style.name.startswith('Heading'):
                    print(f"\n## {text}\n")
                else:
                    print(text)

        # También leer tablas
        print("\n" + "=" * 60)
        print("TABLAS ENCONTRADAS")
        print("=" * 60)

        for i, table in enumerate(doc.tables):
            print(f"\n[TABLA {i+1}]")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_text))

    except ImportError:
        print("⚠️  Necesitas instalar python-docx: pip install python-docx")

if __name__ == '__main__':
    main()
